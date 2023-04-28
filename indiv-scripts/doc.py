#!/usr/bin/python3
# author: suffs811
# Copyright (c) 2023 suffs811
# https://github.com/suffs811/the-terminator.git
# read the README.md file for more details; software distributed under MIT license
# <> purpose: create a .docx file of the .txt report generated by terminator.py
#
# usage: python3 doc.py -r <path_to_report_file>


import os
import sys


parser = argparse.ArgumentParser(description="create a .docx file of the .txt report generated by terminator.py\nusage: python3 doc.py -r <path_to_report_file>")
parser.add_argument("-r", "--report", help="specify path to report file", required=True)
args = parser.parse_args()
report = args.report


# check if docx is installed on machine
def lib_check():
	try:
		import docx
		"docx" in sys.modules
	except:
		return False


# make Word (docx) file and fill with contents from terminator.py output
def doc_make(report):
	# import docx library
	from docx import Document

	# get report name
   	rsplit = report.split("/")
   	fname = rsplit[-1]
   	fsplit = fname.split(".")
   	cut = fsplit[0].strip()

	# create and fill document
	document = Document()
   	r = open("{}".format(report))
   	e = open("/terminator/enum.txt")
   	p = open("/terminator/priv.txt")
   	x = open("/terminator/data_exfil.txt")
   	head = r.readline()
   	ee = e.read()
   	pp = p.read()
   	xx = x.read()

   	document.add_heading(head, 0)
   	document.add_heading("Enumeration", level=1)
   	document.add_paragraph(ee)
   	document.add_page_break()
   	document.add_heading("Exploitation / Initial Shell", level=1)
   	document.add_paragraph("*** ADD YOUR EXPLOITION METHOD FOR THE INITAL SHELL HERE ***")
   	document.add_page_break()
   	document.add_heading("Privilege Escalation", level=1)
   	document.add_paragraph(pp)
   	document.add_page_break()
   	document.add_heading("Persistence and Data Exfiltration", level=1)
   	document.add_paragraph(xx)
   	document.save("{}.docx".format(cut))
   	os.system("mv {}.docx /terminator/{}.docx".format(cut,cut))

   	r.close()
   	e.close()
   	p.close()
   	x.close()

   	print("-+- Word document saved to /terminator/{}.docx -+-".format(cut))


# call functions
lib = lib_check()
if lib:
	doc_make(report)
else:
	print("\n*** 'python-docx' is not installed on your machine; please run 'pip install python-docx' in your terminal ***")
